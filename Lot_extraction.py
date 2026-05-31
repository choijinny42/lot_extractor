from docx import Document
import pandas as pd
import re
from rapidfuzz import fuzz

# synonym dictionary
LOT_SYNONYMS = [
    "Չափաբաժինների համարներով",
    "չափաբաժին",
    "չափաբաժինների համարներով",
    "Չ/Հ",
    "Հրավերով նախատեսված չափաբաժնի համարը",
    "համարները"
]
FUZZ_THRESHOLD = 85
# helper functions

def deduplicate_headers(headers):
    seen = {}
    result = []
    for h in headers:
        if h in seen:
            seen[h] += 1
            result.append(f"{h}_{seen[h]}")
        else:
            seen[h] = 0
            result.append(h)
    return result

def get_element_text(element):
    texts = element._element.xpath(".//w:t")
    return " ".join(t.text for t in texts if t.text)

def normalize(text):
    if not text:
        return ""

    text = text.lower()
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"[^\w\s]", "", text)
    return text.strip()

def get_name(doc, term="Պատվիրատու "):
    results = []
    collecting = False
    buffer = []

    for p in doc.paragraphs:
        text = get_element_text(p).strip()  # use full-run text

        if not collecting and term.lower() in text.lower():
            collecting = True
            after = text.lower().split(term.lower(), 1)[1].strip()
            # get the original-case version
            original = get_element_text(p)
            after_original = original[original.lower().find(term.lower()) + len(term):].strip()
            if after_original:
                buffer.append(after_original)
            continue

        if collecting and not text:
            results.append(" ".join(buffer))
            buffer = []
            collecting = False
            continue

        if collecting:
            buffer.append(text)

    if collecting and buffer:
        results.append(" ".join(buffer))

    return results[0] if results else ""

def safe_filename(text, max_len=80):
    text = text.strip()

    # remove invalid filename chars (Windows)
    text = re.sub(r'[<>:"/\\|?*]', "", text)

    # collapse whitespace
    text = re.sub(r"\s+", " ", text)

    return text[:max_len]

def is_lot_header(header):
    h = normalize(header)
    
    for syn in LOT_SYNONYMS:
        score = fuzz.partial_ratio(h, normalize(syn))
        if score >= FUZZ_THRESHOLD:
            return True
    return False

def cell_has_number(cell):
    # Case 1: Word automatic numbering
    for p in cell.paragraphs:
        if p._p.xpath(".//w:numPr"):
            return True

    # Case 2: actual digit text
    text = get_element_text(cell).strip()
    if re.match(r"^\d+$",text):
        return True

    return False

def get_auto_num_value(cell, table):
    """
    Returns the auto-number integer for a cell using Word's numPr,
    by counting its position among sibling rows that share the same numId.
    """
    for p in cell.paragraphs:
        num_pr = p._p.xpath(".//w:numPr")
        if not num_pr:
            continue
        
        # Get numId and ilvl
        ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        num_id_el = p._p.xpath(".//w:numId")
        ilvl_el = p._p.xpath(".//w:ilvl")
        
        if not num_id_el:
            continue
        
        num_id = num_id_el[0].get(f"{{{ns}}}val")
        ilvl = ilvl_el[0].get(f"{{{ns}}}val") if ilvl_el else "0"
        
        # Count position of this row among rows with the same numId at same level
        count = 0
        for row in table.rows:
            for c in row.cells:
                for cp in c.paragraphs:
                    cp_num_id_el = cp._p.xpath(".//w:numId")
                    cp_ilvl_el = cp._p.xpath(".//w:ilvl")
                    if not cp_num_id_el:
                        continue
                    cp_num_id = cp_num_id_el[0].get(f"{{{ns}}}val")
                    cp_ilvl = cp_ilvl_el[0].get(f"{{{ns}}}val") if cp_ilvl_el else "0"
                    if cp_num_id == num_id and cp_ilvl == ilvl:
                        count += 1
                    if cp._p is p._p:
                        return count
    return None

# data extraction functions
def find_lot_num(doc_path):
    doc = Document(doc_path)

    for table in doc.tables:
        matrix = [[get_element_text(cell).strip() for cell in row.cells] for row in table.rows]
        headers, data_start = get_headers(table)
        #print("HEADERS:", headers)

        lot_col_index = None
        for i, h in enumerate(headers):
            if is_lot_header(h):
                lot_col_index = i
                break

        if lot_col_index is None:
            continue

        lot_nums = []
        for row_idx, row in enumerate(table.rows[data_start:], start=data_start):
            cell = row.cells[lot_col_index]
            text = get_element_text(cell).strip()

            # Try plain text first
            m = re.match(r"\d+", text)
            if m:
                lot_nums.append(m.group())
            else:
                # Fall back to auto-numbering position
                auto_num = get_auto_num_value(cell, table)
                if auto_num is not None:
                    lot_nums.append(str(auto_num))

        if lot_nums:
            return lot_nums, len(lot_nums)

    return [], 0

def get_headers(table):
    matrix = [[get_element_text(cell).strip() for cell in row.cells] for row in table.rows]

    n_cols = len(matrix[0])

    data_start = 1
    for i, row in enumerate(table.rows):
        first_cell_text = get_element_text(row.cells[0]).strip()
        if re.match(r"^\d+$", first_cell_text):
            data_start = i
            break

    headers = []

    for col in range(n_cols):
        parts = []
        last = None

        for row in range(0, data_start):
            value = matrix[row][col]
            if value and value != last:
                parts.append(value)
                last = value

        headers.append(" | ".join(parts))

    headers = [h.split(" | ")[-1] for h in headers]

    return headers, data_start

def clean_by_lot(df, headers):
    # Find lot column using your existing logic
    lot_col = None
    for col in df.columns:
        if is_lot_header(col):
            lot_col = col
            break

    if lot_col is None:
        return df  # no lot column found, return unchanged

    df[lot_col] = df[lot_col].astype(str).str.strip()

    # Remove empty lot rows
    df = df[df[lot_col] != ""]

    # Remove vertically-merged duplicates
    df = df.loc[~df[lot_col].duplicated()]

    return df

def debug_tables(doc_path):
    doc = Document(doc_path)

    for t_index, table in enumerate(doc.tables):
        print(f"\n===== TABLE {t_index} =====")

        matrix = [[get_element_text(cell).strip() for cell in row.cells] for row in table.rows]

        for r_index, row in enumerate(matrix):
            print(f"Row {r_index}: {row}")

def get_cell_text_with_autonumber(cell, table):
    text = get_element_text(cell).strip()
    if text:
        return text
    # Fall back to auto-number position
    auto_num = get_auto_num_value(cell, table)
    if auto_num is not None:
        return str(auto_num)
    return ""

# def get_relevant_tables(doc_path):
#     doc = Document(doc_path)
#     relevant_tables = []

#     for table in doc.tables:
#         matrix = [
#             [get_cell_text_with_autonumber(cell, table) for cell in row.cells]
#             for row in table.rows
#         ]

#         headers, data_start = get_headers(table)

#         # New condition: check header instead of row count
#         if any(is_lot_header(h) for h in headers):
#             relevant_tables.append((matrix, data_start))

#     return relevant_tables

def get_relevant_tables(doc_path):
    doc = Document(doc_path)
    relevant_tables = []
    print(f"get_relevant_tables: scanning {len(doc.tables)} tables")

    for t_idx, table in enumerate(doc.tables):
        matrix = [
            [get_cell_text_with_autonumber(cell, table) for cell in row.cells]
            for row in table.rows
        ]
        headers, data_start = get_headers(table)
        print(f"  table {t_idx}: data_start={data_start}, headers={headers!r}")

        matched = [(h, is_lot_header(h)) for h in headers]
        print(f"    is_lot_header per col: {matched}")

        if any(is_lot_header(h) for h in headers):
            relevant_tables.append((matrix, data_start))

    print(f"  -> {len(relevant_tables)} relevant tables")
    return relevant_tables

# def save_relevant_tables(doc_path, base_name):
#     relevant_tables = get_relevant_tables(doc_path)

#     for i, (matrix, data_start) in enumerate(relevant_tables):
#         header = matrix[data_start - 1]
#         data = matrix[data_start:]

#         df = pd.DataFrame(data, columns=header)
#         df.to_csv(f"{base_name}_{i}.csv", index=False, encoding="utf-8-sig")

def save_relevant_tables(doc_path, base_name):
    relevant_tables = get_relevant_tables(doc_path)
    dataframes = []

    for i, (matrix, data_start) in enumerate(relevant_tables):
        data = matrix[data_start:]
        header = deduplicate_headers(matrix[data_start - 1])
        df = pd.DataFrame(data, columns=header)

        # CLEAN HERE
        df = clean_by_lot(df, header)

        dataframes.append(df)

    return dataframes

