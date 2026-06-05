from docx import Document
import pandas as pd
import re
# use this program if there is a second document with table info

def deduplicate_headers(headers):
    seen = {}
    result = []
    for h in headers:
        if h in seen:
            seen[h] += 1
            result.append(f"{h}_{seen[h]}")
        else:
            seen[h] = 0
            result.append(h)
    return result

def get_cell_text(cell):
    texts = cell._element.xpath(".//w:t")
    return " ".join(t.text for t in texts if t.text)

def starts_with_number(text):
    if not text:
        return False
    return bool(re.match(r"^\s*\d+[\.\)]?", text))

def paragraph_is_numbered(paragraph):
    return bool(paragraph._p.xpath(".//w:numPr"))

def cell_has_number(cell):
    # Case 1 — Word numbering
    for p in cell.paragraphs:
        if paragraph_is_numbered(p):
            return True

    # Case 2 — digit numbering text
    text = get_cell_text(cell).strip()
    if starts_with_number(text):
        return True

    return False

def find_data_start(table):
    for i, row in enumerate(table.rows):
        if any(cell_has_number(cell) for cell in row.cells):
            return i
    return 0


def get_cell_text(cell):
    texts = cell._element.xpath(".//w:t")
    return " ".join(t.text for t in texts if t.text)

def word_to_csv(doc_path, new_columns, table_name):
    doc = Document(doc_path)

    for t_index, table in enumerate(doc.tables):
        data = []

        for r_index, row in enumerate(table.rows):
            if r_index <= 1:
                continue

            text = [get_cell_text(cell).strip() for cell in row.cells]
            data.append(text)

        df = pd.DataFrame(data, columns=new_columns)

        df.to_csv(f"{table_name}_{t_index}.csv", index=False, encoding="utf-8-sig")

    return df


def get_table(doc_path):
    doc = Document(doc_path)
    return doc.tables
    
def get_headers(table):
    matrix = [[get_cell_text(cell).strip() for cell in row.cells] for row in table.rows]

    if not matrix:
        return [], 0

    n_cols = len(matrix[0])

    # detect where data starts
    data_start = find_data_start(table)

    # build headers from rows above data_start
    headers = []

    for col in range(n_cols):
        parts = []
        last = None

        for row in range(0, data_start):
            value = matrix[row][col]

            if value and value != last:
                parts.append(value)
                last = value

        header = " | ".join(parts) if parts else f"col_{col}"
        headers.append(header)

    # keep last header level 
    headers = [h.split(" | ")[-1] for h in headers]

    return headers, data_start




def save_csv(doc_path):
    doc = Document(doc_path)
    
    dataframes = []
    
    for t_index, table in enumerate(doc.tables):
        headers, data_start_row = get_headers(table)
        headers = deduplicate_headers(headers)

        matrix = [[get_cell_text(cell).strip() for cell in row.cells] for row in table.rows]

        data = matrix[data_start_row:]

        df = pd.DataFrame(data, columns=headers)
        if df.empty:
            continue
        
        dataframes.append(df)
    
    return dataframes

